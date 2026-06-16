from io import BytesIO
from typing import Optional

from app.config import RULES_DIR
from app.rules.loader import RuleLoader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt


def _set_run_font(run, size: float, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def _format_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(28)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def _add_multiline_paragraphs(document: Document, content: str) -> None:
    for block in content.split("\n\n"):
        lines = [line for line in block.split("\n") if line.strip()]
        if not lines:
            continue
        paragraph = document.add_paragraph()
        _format_body_paragraph(paragraph)
        first_run = paragraph.add_run(lines[0])
        _set_run_font(first_run, 12)
        for line in lines[1:]:
            next_run = paragraph.add_run("\n" + line)
            _set_run_font(next_run, 12)


def _add_centered_paragraph(document: Document, text: str, size: float, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    _set_run_font(run, size, bold=bold)


def _add_labeled_centered_paragraph(
    document: Document, label: str, value: str, size: float
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = paragraph.add_run(f"{label}：")
    _set_run_font(label_run, size, bold=True)
    value_run = paragraph.add_run(value)
    _set_run_font(value_run, size)


def _add_heading_paragraph(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    if level == 1:
        _set_run_font(run, 16, bold=True)
    else:
        _set_run_font(run, 14, bold=True)


def _add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def _configure_document(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    for section in document.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(90)
        section.right_margin = Pt(72)

        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run("突发环境事件应急预案送审稿")
        _set_run_font(header_run, 10)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("供专家评审使用")
        _set_run_font(footer_run, 10)


def _add_cover_page(document: Document, export_payload: dict[str, str]) -> None:
    for _ in range(4):
        document.add_paragraph()
    _add_centered_paragraph(document, export_payload["cover_title"], 22, bold=True)
    document.add_paragraph()
    if export_payload.get("project_name"):
        _add_centered_paragraph(document, export_payload["project_name"], 16, bold=True)
    for _ in range(8):
        document.add_paragraph()
    _add_centered_paragraph(document, "送审稿", 16, bold=True)
    for _ in range(3):
        document.add_paragraph()
    _add_labeled_centered_paragraph(
        document,
        "编制单位",
        export_payload.get("project_name", "项目单位"),
        14,
    )
    _add_labeled_centered_paragraph(
        document,
        "编制日期",
        export_payload.get("plan_issue_date", "拟定实施日期"),
        14,
    )


def _build_numbered_titles(titles: list[str]) -> list[str]:
    return [f"{index}. {title}" for index, title in enumerate(titles, start=1)]


def _build_numbered_body(body: str) -> str:
    blocks = [block for block in body.split("\n\n") if block.strip()]
    numbered_blocks: list[str] = []
    for index, block in enumerate(blocks, start=1):
        lines = [line for line in block.split("\n") if line.strip()]
        if not lines:
            continue
        numbered_title = f"{index}. {lines[0]}"
        if len(lines) > 1:
            numbered_blocks.append("\n".join([numbered_title, *lines[1:]]))
        else:
            numbered_blocks.append(numbered_title)
    return "\n\n".join(numbered_blocks)


def _build_full_preview(
    project_name: str,
    ordered_titles: list[str],
    numbered_titles: list[str],
    body: str,
    numbered_body: str,
    preface_payload: Optional[dict[str, str]],
    attachments_payload: Optional[dict[str, object]],
) -> str:
    full_preview_parts = [
        "封面",
        "突发环境事件应急预案",
        project_name,
    ]
    if preface_payload:
        full_preview_parts.extend(
            [
                preface_payload["filing_form"],
                preface_payload["filing_directory"],
                preface_payload["release_order"],
                preface_payload["compilation_notes"],
                "正文目录",
                "\n".join(numbered_titles),
                numbered_body,
            ]
        )
    else:
        full_preview_parts.extend(["正文目录", "\n".join(numbered_titles), numbered_body])
    if attachments_payload:
        full_preview_parts.extend(
            [
                attachments_payload["appendix_catalog"],
                attachments_payload["communication_list"],
                attachments_payload["materials_list"],
            ]
        )
    return "\n\n".join(full_preview_parts)


def _add_toc_block(document: Document, toc_text: str) -> None:
    _add_heading_paragraph(document, "正文目录", level=1)
    for line in [line for line in toc_text.split("\n") if line.strip()]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(line)
        _set_run_font(run, 12)


def _add_section_blocks(document: Document, body: str) -> None:
    for block in body.split("\n\n"):
        lines = [line for line in block.split("\n") if line.strip()]
        if not lines:
            continue
        title = lines[0]
        content_lines = lines[1:]
        _add_heading_paragraph(document, title, level=1)
        for content in content_lines:
            paragraph = document.add_paragraph()
            _format_body_paragraph(paragraph)
            run = paragraph.add_run(content)
            _set_run_font(run, 12)


def render_filing_directory() -> str:
    loader = RuleLoader(RULES_DIR)
    data = loader.load_json("filing_directory_template.json")
    items = data["items"]
    return "\n".join(f"{index}. {item}" for index, item in enumerate(items, start=1))


def render_filing_form(facts: dict[str, str]) -> str:
    loader = RuleLoader(RULES_DIR)
    data = loader.load_json("filing_form_template.json")
    lines = [data["title"]]
    for field in data["fields"]:
        value = facts.get(field["key"], field["fallback"])
        lines.append(f"{field['label']}：{value}")
    return "\n".join(lines)


def render_release_order(facts: dict[str, str]) -> str:
    loader = RuleLoader(RULES_DIR)
    data = loader.load_json("release_order_template.json")
    org_type = facts.get("organization_type", "default")
    org_label = data["organization_labels"].get(org_type, data["organization_labels"]["default"])
    org_unit = data["organization_units"].get(org_type, data["organization_units"]["default"])
    lines = [
        line.format(
            org_label=org_label,
            org_unit=org_unit,
            company_name=facts.get("company_name", "本单位"),
            plan_name=facts.get("plan_name", f"{facts.get('company_name', '本单位')}突发环境事件应急预案"),
            plan_issue_date=facts.get("plan_issue_date", "待补充日期"),
            legal_representative=facts.get("legal_representative", "待补充签署人"),
        )
        for line in data["template"]
    ]
    return "\n".join(lines)


def render_compilation_notes(facts: dict[str, str]) -> str:
    loader = RuleLoader(RULES_DIR)
    data = loader.load_json("compilation_notes_template.json")
    lines: list[str] = ["编制说明"]
    for section in data["sections"]:
        lines.append(section["title"])
        lines.append(section["template"].format(**facts))
    return "\n".join(lines)


def render_export_payload(
    project_name: str,
    sections: dict[str, str],
    preface_payload: Optional[dict[str, str]] = None,
    plan_issue_date: str = "",
    attachments_payload: Optional[dict[str, object]] = None,
) -> dict[str, str]:
    ordered_titles = list(sections.keys())
    numbered_titles = _build_numbered_titles(ordered_titles)
    body_parts = [f"{title}\n{sections[title]}" for title in ordered_titles]
    body = "\n\n".join(body_parts)
    numbered_body = _build_numbered_body(body)
    return {
        "cover_title": "突发环境事件应急预案",
        "project_name": project_name,
        "plan_issue_date": plan_issue_date,
        "table_of_contents": "\n".join(numbered_titles),
        "body": body,
        "filing_form": preface_payload["filing_form"] if preface_payload else "",
        "filing_directory": preface_payload["filing_directory"] if preface_payload else "",
        "release_order": preface_payload["release_order"] if preface_payload else "",
        "compilation_notes": preface_payload["compilation_notes"] if preface_payload else "",
        "full_preview": _build_full_preview(
            project_name=project_name,
            ordered_titles=ordered_titles,
            numbered_titles=numbered_titles,
            body=body,
            numbered_body=numbered_body,
            preface_payload=preface_payload,
            attachments_payload=attachments_payload,
        ),
        "appendix_catalog": str(attachments_payload["appendix_catalog"]) if attachments_payload else "",
        "communication_list": str(attachments_payload["communication_list"]) if attachments_payload else "",
        "materials_list": str(attachments_payload["materials_list"]) if attachments_payload else "",
    }


def build_word_document_bytes(export_payload: dict[str, str]) -> bytes:
    document = Document()
    _configure_document(document)
    _add_cover_page(document, export_payload)
    preface_blocks = [
        export_payload.get("filing_form", ""),
        export_payload.get("filing_directory", ""),
        export_payload.get("release_order", ""),
        export_payload.get("compilation_notes", ""),
    ]
    preface_blocks = [block for block in preface_blocks if block]
    for block in preface_blocks:
        _add_page_break(document)
        lines = [line for line in block.split("\n") if line.strip()]
        if not lines:
            continue
        _add_heading_paragraph(document, lines[0], level=1)
        content = "\n".join(lines[1:]) if len(lines) > 1 else ""
        if content:
            _add_multiline_paragraphs(document, content)

    _add_page_break(document)
    if export_payload.get("table_of_contents"):
        _add_toc_block(document, export_payload["table_of_contents"])
    _add_page_break(document)
    _add_section_blocks(document, export_payload["body"])

    attachment_blocks = [
        export_payload.get("appendix_catalog", ""),
        export_payload.get("communication_list", ""),
        export_payload.get("materials_list", ""),
    ]
    for block in [block for block in attachment_blocks if block]:
        _add_page_break(document)
        lines = [line for line in block.split("\n") if line.strip()]
        if not lines:
            continue
        _add_heading_paragraph(document, lines[0], level=1)
        content = "\n".join(lines[1:]) if len(lines) > 1 else ""
        if content:
            _add_multiline_paragraphs(document, content)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
