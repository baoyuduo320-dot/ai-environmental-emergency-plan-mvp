from io import BytesIO

from docx import Document

from app.export.service import (
    build_word_document_bytes,
    render_compilation_notes,
    render_export_payload,
    render_filing_directory,
    render_filing_form,
    render_release_order,
)


def test_render_export_payload_contains_cover_and_sections():
    payload = render_export_payload(
        project_name="苏州示例化工有限公司",
        sections={
            "总则": "总则内容",
            "企业基本情况": "企业基本情况内容",
        },
        preface_payload={
            "filing_form": "备案表内容",
            "filing_directory": "备案目录内容",
            "release_order": "发布令内容",
            "compilation_notes": "编制说明内容",
        },
        plan_issue_date="2026年6月15日",
    )
    assert payload["cover_title"] == "突发环境事件应急预案"
    assert "总则内容" in payload["body"]
    assert "封面" in payload["full_preview"]
    assert "备案表内容" in payload["full_preview"]
    assert "编制说明内容" in payload["full_preview"]
    assert "正文目录" in payload["full_preview"]
    assert "1. 总则" in payload["full_preview"]
    assert payload["table_of_contents"].split("\n")[0] == "1. 总则"


def test_render_release_order_uses_template_values():
    payload = render_release_order(
        {
            "company_name": "苏州示例化工有限公司",
            "plan_name": "苏州示例化工有限公司突发环境事件应急预案",
            "plan_issue_date": "2026年6月15日",
            "legal_representative": "张三",
        }
    )
    assert "公司各部门：" in payload
    assert "苏州示例化工有限公司突发环境事件应急预案" in payload
    assert "签署人：张三" in payload


def test_render_filing_directory_contains_fixed_items():
    payload = render_filing_directory()
    assert "1. 突发环境事件应急预案备案表" in payload
    assert "5. 环境应急预案评审意见" in payload


def test_render_filing_form_contains_structured_fields():
    payload = render_filing_form(
        {
            "company_name": "苏州示例化工有限公司",
            "organization_code": "91320500MA12345678",
            "legal_representative": "张三",
            "contact_person": "李四",
            "contact_phone": "13800000000",
            "address": "苏州市工业园区示范路88号",
            "longitude_latitude": "东经120°00′00″，北纬31°00′00″",
            "plan_name": "苏州示例化工有限公司突发环境事件应急预案",
            "risk_level": "一般",
            "plan_issue_date": "2026年6月15日",
        }
    )
    assert "企业事业单位突发环境事件应急预案备案表" in payload
    assert "单位名称：苏州示例化工有限公司" in payload
    assert "机构代码：91320500MA12345678" in payload
    assert "法定代表人：张三" in payload
    assert "联系人：李四" in payload
    assert "联系人电话：13800000000" in payload
    assert "地址：苏州市工业园区示范路88号" in payload
    assert "经纬度：东经120°00′00″，北纬31°00′00″" in payload
    assert "风险级别：一般" in payload
    assert "实施日期：2026年6月15日" in payload


def test_render_compilation_notes_uses_review_ready_wording():
    payload = render_compilation_notes(
        {
            "company_name": "苏州示例化工有限公司",
            "plan_name": "苏州示例化工有限公司突发环境事件应急预案",
            "region": "苏州市工业园区",
            "industry": "精细化工",
            "risk_level": "一般",
            "risk_sources": "危化品储罐区、危废暂存间和重点生产装置",
            "risk_materials": "危险化学品、危险废物和辅料的储量、位置和危险特性",
            "pollutant_discharge": "废气治理设施、污水处理站和事故状态下截流收集措施",
            "environmental_receptors": "周边居民点、雨水排口和地表水体",
            "emergency_org_structure": "综合协调、现场处置、监测预警和物资保障工作组",
            "emergency_resources": "防化服、堵漏工具、吸附材料和监测仪器",
            "response_process": "首报、续报、终报和联动处置流程",
        }
    )
    assert "编制说明" in payload
    assert "形成了可供进一步论证完善的阶段性成果" in payload
    assert "为后续结合论证意见进一步优化完善提供依据" in payload
    assert "已完成评审" not in payload


def test_build_word_document_bytes_contains_core_sections():
    payload = render_export_payload(
        project_name="苏州示例化工有限公司",
        sections={
            "总则": "总则内容",
            "培训与演练": "培训与演练内容",
        },
        preface_payload={
            "filing_form": "备案表内容",
            "filing_directory": "备案目录内容",
            "release_order": "发布令内容",
            "compilation_notes": "编制说明内容",
        },
        plan_issue_date="2026年6月15日",
    )
    document_bytes = build_word_document_bytes(payload)
    assert document_bytes.startswith(b"PK")

    document = Document(BytesIO(document_bytes))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    assert "突发环境事件应急预案" in text
    assert "苏州示例化工有限公司" in text
    assert "2026年6月15日" in text
    assert "编制单位：苏州示例化工有限公司" in text
    assert "编制日期：2026年6月15日" in text
    assert "备案表内容" in text
    assert "编制说明内容" in text
    assert "培训与演练内容" in text
    assert paragraphs[0] == "突发环境事件应急预案"
    assert "正文目录" in paragraphs
    assert "1. 总则" in paragraphs
    assert "2. 培训与演练" in paragraphs
    header_text = "\n".join(paragraph.text for paragraph in document.sections[0].header.paragraphs)
    footer_text = "\n".join(paragraph.text for paragraph in document.sections[0].footer.paragraphs)
    assert "突发环境事件应急预案送审稿" in header_text
    assert "供专家评审使用" in footer_text
