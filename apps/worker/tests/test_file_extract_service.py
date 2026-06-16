from io import BytesIO

from docx import Document

from app.file_extract.service import clean_extracted_text, extract_text_from_upload


def test_extract_text_from_docx_upload_reads_text():
    document = Document()
    document.add_paragraph("企业名称：测试公司")
    document.add_paragraph("所属行业：精细化工")
    buffer = BytesIO()
    document.save(buffer)

    result = extract_text_from_upload("sample.docx", buffer.getvalue())
    assert "企业名称：测试公司" in result["text"]
    assert result["warnings"] == []


def test_extract_text_from_pdf_upload_reads_literal_text():
    pdf_bytes = (
        "%PDF-1.4\n"
        "1 0 obj\n<< /Length 35 >>\nstream\n"
        "BT\n"
        "(企业名称：测试公司) Tj\n"
        "ET\n"
        "endstream\n"
        "endobj\n"
        "%%EOF"
    ).encode("utf-8")

    result = extract_text_from_upload("sample.pdf", pdf_bytes)
    assert "企业名称：测试公司" in result["text"]


def test_extract_text_from_unsupported_upload_raises_error():
    try:
        extract_text_from_upload("sample.xlsx", b"demo")
    except ValueError as error:
        assert "暂不支持" in str(error)
    else:
        raise AssertionError("expected ValueError for unsupported file type")


def test_clean_extracted_text_removes_noise_lines():
    cleaned = clean_extracted_text(
        "封面\n\n企业名称：测试公司\n1\n\n企业名称：测试公司\n正文目录\n所属行业：精细化工\n"
    )
    assert "封面" not in cleaned
    assert "正文目录" not in cleaned
    assert "\n1\n" not in cleaned
    assert "企业名称：测试公司" in cleaned
